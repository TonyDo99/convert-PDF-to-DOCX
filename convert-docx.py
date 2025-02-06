# import os
# from pdf2docx import parse

# # Define paths for PDF and DOCX files
# pdfs = os.path.join(os.getcwd(), 'pdfs')  # Use os.path.join() for safe paths
# docx_dir = os.path.join(os.getcwd(), 'docxs')

# # Ensure docx directory exists
# if not os.path.exists(docx_dir):
#     os.mkdir(docx_dir)

# count = 1

# # Loop through the files in the pdfs directory
# for page in os.listdir(pdfs):
#     # Check if the file is a PDF
#     if page.endswith('.pdf'):
#         pdf_path = os.path.join(pdfs, f"file-{count}.pdf")  # Full PDF path
#         docx_path = os.path.join(docx_dir, f"file-{count}.docx")  # Full DOCX path
        
#         print(f"Converting {pdf_path} to {docx_path}")
        
#         # Convert the PDF to DOCX
#         # convert pdf to docx
        
#         # cv = Converter(pdf_path)
#         # cv.convert(docx_path)
#         # cv.close()
#         parse(pdf_path, docx_path)
#         count += 1

# print("All PDFs have been converted!")


import os
import fitz  # PyMuPDF
from pdf2docx import Converter
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configurations
INPUT_DOCX = "document.docx"
OUTPUT_DIR = "docxs"
N_PAGES = 1  # Adjust as needed

# Ensure output directory exists
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Step 1: Convert DOCX to PDF using LibreOffice
pdf_path = INPUT_DOCX.replace(".docx", ".pdf")
os.system(f"libreoffice --headless --convert-to pdf {INPUT_DOCX}")

# Step 2: Extract text per page from PDF
pdf_doc = fitz.open(pdf_path)
num_pages = len(pdf_doc)

# Step 3: Split DOCX based on pages
for start_page in range(0, num_pages, N_PAGES):
    end_page = start_page + N_PAGES
    
    # Convert page range back to DOCX
    output_docx = os.path.join(OUTPUT_DIR, f"file-{start_page + N_PAGES}.docx")
    cv = Converter(pdf_path)
    cv.convert(output_docx, start=start_page, end=end_page)
    cv.close()

    # Post-process the DOCX to fix table styles
    doc = Document(output_docx)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Ensure consistent cell margins and borders
                cell.paragraphs[0].paragraph_format.space_before = 0
                cell.paragraphs[0].paragraph_format.space_after = 0
                cell.paragraphs[0].paragraph_format.line_spacing = 1.0
            
        # Adjust column widths (example for a 3-column table)
        if len(table.columns) == 3:
            table.columns[0].width = Inches(0.5)
            table.columns[1].width = Inches(4.5)
    doc.save(output_docx)
    print(f"Saved: {output_docx}")

print("Splitting completed.")